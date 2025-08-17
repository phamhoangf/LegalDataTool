import os
import mimetypes
import chardet
from typing import Dict, Tuple

def detect_encoding(file_content: bytes) -> str:
    """
    Phát hiện encoding của file
    """
    try:
        detected = chardet.detect(file_content)
        encoding = detected.get('encoding', 'utf-8')
        confidence = detected.get('confidence', 0)
        
        # Nếu confidence thấp hoặc không detect được, thử các encoding phổ biến
        if confidence < 0.7 or not encoding:
            for enc in ['utf-8', 'cp1252', 'latin1', 'utf-16']:
                try:
                    file_content.decode(enc)
                    return enc
                except UnicodeDecodeError:
                    continue
            return 'utf-8'  # fallback
        
        return encoding
    except:
        return 'utf-8'

def process_text_file(file_content: bytes, filename: str = None) -> Dict:
    """
    Xử lý file text với auto-detect encoding
    """
    try:
        encoding = detect_encoding(file_content)
        
        # Thử decode với encoding đã detect
        try:
            content = file_content.decode(encoding)
        except UnicodeDecodeError:
            # Fallback: thử các encoding khác
            for enc in ['utf-8', 'cp1252', 'latin1', 'utf-16']:
                try:
                    content = file_content.decode(enc)
                    encoding = enc
                    break
                except UnicodeDecodeError:
                    continue
            else:
                # Nếu tất cả đều fail, decode với errors='ignore'
                content = file_content.decode('utf-8', errors='ignore')
                encoding = 'utf-8 (with errors ignored)'
        
        return {
            'success': True,
            'content': content,
            'encoding': encoding,
            'file_type': 'text'
        }
    except Exception as e:
        return {
            'success': False,
            'error': f'Không thể xử lý file text: {str(e)}'
        }

def process_pdf_file(file_content: bytes, filename: str = None) -> Dict:
    """
    Xử lý file PDF bằng PyPDF2
    """
    try:
        import PyPDF2
        import io
        
        pdf_stream = io.BytesIO(file_content)
        reader = PyPDF2.PdfReader(pdf_stream)
        
        content_parts = []
        for page_num, page in enumerate(reader.pages, 1):
            text = page.extract_text()
            if text.strip():
                content_parts.append(f"[Trang {page_num}]\n{text.strip()}")
        
        if not content_parts:
            return {
                'success': False,
                'error': 'Không thể trích xuất text từ PDF. File có thể là PDF ảnh hoặc bị mã hóa.'
            }
        
        # Lấy metadata
        metadata = {}
        if reader.metadata:
            metadata = {
                'title': reader.metadata.get('/Title', ''),
                'author': reader.metadata.get('/Author', ''),
                'subject': reader.metadata.get('/Subject', ''),
                'creator': reader.metadata.get('/Creator', ''),
                'pages': len(reader.pages)
            }
        
        return {
            'success': True,
            'content': '\n\n'.join(content_parts),
            'file_type': 'pdf',
            'metadata': metadata
        }
        
    except ImportError:
        return {
            'success': False,
            'error': 'PyPDF2 chưa được cài đặt. Chạy: pip install PyPDF2'
        }
    except Exception as e:
        return {
            'success': False,
            'error': f'Lỗi xử lý PDF: {str(e)}'
        }

def process_docx_file(file_content: bytes, filename: str = None) -> Dict:
    """
    Xử lý file DOCX bằng python-docx
    """
    try:
        from docx import Document
        import io
        
        doc_stream = io.BytesIO(file_content)
        doc = Document(doc_stream)
        
        content_parts = []
        
        # Trích xuất paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                content_parts.append(para.text.strip())
        
        # Trích xuất tables
        for table in doc.tables:
            table_content = []
            for row in table.rows:
                row_content = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_content.append(cell.text.strip())
                if row_content:
                    table_content.append(' | '.join(row_content))
            
            if table_content:
                content_parts.append('\n[Bảng]\n' + '\n'.join(table_content))
        
        if not content_parts:
            return {
                'success': False,
                'error': 'Không có nội dung text trong file Word'
            }
        
        # Metadata
        metadata = {
            'paragraphs': len(doc.paragraphs),
            'tables': len(doc.tables)
        }
        
        # Document properties
        try:
            core_props = doc.core_properties
            if core_props:
                metadata.update({
                    'title': core_props.title or '',
                    'author': core_props.author or '',
                    'subject': core_props.subject or ''
                })
        except:
            pass
        
        return {
            'success': True,
            'content': '\n\n'.join(content_parts),
            'file_type': 'docx',
            'metadata': metadata
        }
        
    except ImportError:
        return {
            'success': False,
            'error': 'python-docx chưa được cài đặt. Chạy: pip install python-docx'
        }
    except Exception as e:
        return {
            'success': False,
            'error': f'Lỗi xử lý DOCX: {str(e)}'
        }

def get_file_type(filename: str) -> str:
    """
    Xác định loại file từ extension
    """
    if not filename:
        return 'unknown'
    
    _, ext = os.path.splitext(filename.lower())
    
    if ext == '.pdf':
        return 'pdf'
    elif ext in ['.docx', '.doc']:
        return 'docx'
    elif ext in ['.txt', '.text']:
        return 'text'
    elif ext in ['.html', '.htm']:
        return 'html'
    else:
        return 'text'  # Fallback: coi như text file

def process_file(file_content: bytes, filename: str = None) -> Dict:
    """
    Xử lý file dựa trên loại file
    """
    file_type = get_file_type(filename)
    
    if file_type == 'pdf':
        return process_pdf_file(file_content, filename)
    elif file_type == 'docx':
        return process_docx_file(file_content, filename)
    else:
        # Mặc định xử lý như text file
        return process_text_file(file_content, filename)

def validate_file_size(file_content: bytes, max_size_mb: int = 50) -> Tuple[bool, str]:
    """
    Kiểm tra kích thước file
    """
    size_mb = len(file_content) / (1024 * 1024)
    
    if size_mb > max_size_mb:
        return False, f'File quá lớn ({size_mb:.1f}MB). Tối đa {max_size_mb}MB'
    
    return True, ''

def get_supported_formats() -> Dict:
    """
    Trả về danh sách format được hỗ trợ
    """
    formats = {
        'text': {
            'extensions': ['.txt', '.text'],
            'description': 'Plain text files with auto-encoding detection'
        },
        'html': {
            'extensions': ['.html', '.htm'],
            'description': 'HTML files (processed as text)'
        }
    }
    
    # Kiểm tra PyPDF2
    try:
        import PyPDF2
        formats['pdf'] = {
            'extensions': ['.pdf'],
            'description': 'PDF files with text extraction'
        }
    except ImportError:
        pass
    
    # Kiểm tra python-docx
    try:
        from docx import Document
        formats['docx'] = {
            'extensions': ['.docx', '.doc'],
            'description': 'Microsoft Word documents'
        }
    except ImportError:
        pass
    
    return formats
